from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

def sf(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    if level == 1:
        sf(r, 14, bold=True, color=(26, 26, 46))
    elif level == 2:
        sf(r, 12, bold=True, color=(22, 33, 62))
    else:
        sf(r, 11, bold=True, color=(50, 50, 80))

def placeholder(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'[ {text} ]')
    sf(r, 10, italic=True, color=(150, 150, 150))

def citation(code, ref, cite_for, argument, quality='High'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(f'[{code}]  ')
    sf(r1, 9, bold=True, color=(39, 103, 73))
    r2 = p.add_run(f'Quality: {quality}')
    sf(r2, 8, italic=True, color=(120, 120, 120))

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.9)
    p2.paragraph_format.first_line_indent = Cm(-0.4)
    p2.paragraph_format.space_after = Pt(1)
    sf(p2.add_run(ref), 9.5)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.9)
    p3.paragraph_format.space_after = Pt(1)
    sf(p3.add_run('Cite for: '), 9, bold=True)
    sf(p3.add_run(cite_for), 9)

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Cm(0.9)
    p4.paragraph_format.space_after = Pt(9)
    sf(p4.add_run('Argument: '), 9, bold=True)
    sf(p4.add_run(argument), 9, italic=True)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    sf(p.add_run('_' * 90), 6, color=(200, 200, 200))

# ---- TITLE ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
sf(p.add_run('Anxiety in Children and Adolescents: A Systemic Disorder\n'), 18, bold=True, color=(26, 26, 46))
sf(p.add_run('Evidence-Based Approaches and Interventions'), 13, italic=True, color=(80, 80, 100))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p2.add_run('MSc NPMH - Module 4 Assessment B | 2,000 words (+/- 10%) | Harvard WMS Referencing'), 10, color=(120, 120, 120))

doc.add_page_break()

# ---- THESIS ----
heading('Thesis Statement', level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
sf(p.add_run(
    '"Anxiety in children and adolescents is not an individual disorder but a systemic one: '
    'sustained by family dynamics that accommodate and transmit anxiety, and amplified by a digital '
    'and societal environment that exploits adolescent neurobiological vulnerability. Effective, '
    'evidence-based intervention must therefore move beyond the individual child -- engaging family '
    'systems through inclusive CBT and systemic approaches, while addressing the external forces that '
    'generate and maintain anxiety in young people\'s everyday lives."'
), 10.5, italic=True, color=(50, 80, 60))

divider()

# ---- 1. INTRO ----
heading('1. Introduction  (~200 words)', level=1)
placeholder('Hook: anxiety as most common CAMH disorder -- 6.5% globally, 1 in 12 in US (M4 W4 P51)')
placeholder('Undertreatment gap: majority of affected children do not access evidence-based care')
placeholder('Individual framing is insufficient -- introduce ecosystem argument')
placeholder('Roadmap: essay focuses on family system and digital/societal environment')
placeholder('Thesis stated')
divider()

# ---- 2. FAMILY ----
heading('2. The Family System  (~550 words)', level=1)
heading('2.1  Parental Anxiety Transmission', level=2)
placeholder('Social referencing: anxious parents model threat-focused worldviews (M4 W1 curriculum)')
placeholder('Causal (non-genetic) transmission -- central argument for systemic intervention')

citation('PT-1',
    'Ahmadzadeh, Y., Schoeler, T., Han, M. et al. (2021) Systematic Review and Meta-Analysis of '
    'Genetically Informed Research: Associations Between Parent Anxiety and Offspring Internalizing '
    'Problems. Journal of the American Academy of Child and Adolescent Psychiatry, 60(8), pp. 959-972.',
    'Proving causal (non-genetic) parental anxiety transmission (r=.13, quasi-experimental, n>12,700).',
    'Postnatal parental anxiety causally -- not merely genetically -- transmits to children, making '
    'parental intervention a legitimate and necessary treatment target.',
    quality='High')

citation('PT-2',
    'Zecchinato, F., Ahmadzadeh, Y., Kreppner, J. and Lawrence, P.J. (2024) Paternal Anxiety and '
    'Emotional and Behavioral Outcomes in Offspring: A Systematic Review and Meta-Analysis. Journal '
    'of the American Academy of Child and Adolescent Psychiatry, 63(11), pp. 1142-1156.',
    'Extending transmission argument to fathers (98 studies, n~55,000; r=.13-.15).',
    'Both maternal and paternal anxiety contribute to child anxiety, reinforcing whole-family '
    'intervention over mother-focused approaches.',
    quality='High')

citation('PT-3',
    'Bennett, S., McClure, G. and Yap, M. (2025) Differences in Parental Factors Between Parents '
    'With and Without Depression or Anxiety: A Systematic Review and Meta-Analysis. Journal of '
    'Affective Disorders, 369, pp. 120085.',
    'Identifying the behavioural mechanism: anxious parents show reduced warmth and less effective discipline.',
    'Parental anxiety transmits via modifiable parenting behaviours -- reduced warmth, overcontrol -- '
    'which are directly addressable through family intervention.',
    quality='High')

heading('2.2  Family Accommodation', level=2)
placeholder('Define family accommodation (FA): parents adjusting routines to reduce child distress')
placeholder('FA maintains and reinforces avoidance cycle -- sustaining rather than resolving anxiety')

citation('FA-2',
    'Fox, R. and Fleming, L. (2025) Parental Cognitions and Child Anxiety: A Systematic Review. '
    'Journal of Anxiety Disorders, 101, pp. 103021.',
    'Explaining the cognitive mechanism linking parental anxiety to accommodation behaviour (31 studies).',
    'Negative parental cognitions mediate the path from parental anxiety to accommodation -- '
    'identifying a clear cognitive intervention target upstream of the behaviour.',
    quality='High')

heading('2.3  Family-Inclusive Interventions', level=2)
placeholder('CBT as gold standard -- family-inclusive format strengthens outcomes (M4 W4 P56, curriculum)')

citation('CBT-2',
    'Wang, Z., Whiteside, S.P.H., Sim, L. et al. (2017) Comparative Effectiveness and Safety of '
    'Cognitive Behavioral Therapy and Pharmacotherapy for Childhood Anxiety Disorders: A Systematic '
    'Review and Meta-Analysis. JAMA Pediatrics, 171(11), pp. 1049-1056.',
    'Establishing CBT as gold standard; combination CBT + SSRI > either alone (115 studies, n=7,719).',
    'CBT is the most robustly evidenced treatment for childhood anxiety, with family involvement '
    'enhancing outcomes -- the empirical foundation for a systemic treatment approach.',
    quality='High')

citation('SPACE-1',
    'Lebowitz, E.R., Marin, C., Martino, A. et al. (2019) Parent-Based Treatment as Efficacious as '
    'Cognitive Behavioral Therapy for Childhood Anxiety: A Randomized Noninferiority Study of SPACE. '
    'Journal of the American Academy of Child and Adolescent Psychiatry, 58(12), pp. 1170-1181.',
    'Demonstrating parent-only SPACE treatment equals child CBT outcomes; achieves greater FA reduction.',
    'SPACE demonstrates that treating parents alone -- without direct child contact -- produces '
    'equivalent anxiety reduction, fundamentally challenging the individual-child treatment model.',
    quality='High')

placeholder('Critical evaluation: adolescent autonomy limits family involvement -- agreed at outset and reviewed (M4 W4 P56)')
divider()

# ---- 3. DIGITAL ----
heading('3. The Digital and Societal Environment  (~550 words)', level=1)
heading('3.1  Neurobiological Vulnerability in Adolescence', level=2)
placeholder('vmPFC development during puberty: heightened social self-evaluation and social anxiety risk (M4 W1 P06 -- must cite)')

citation('NB-1',
    'Casey, B.J., Glatt, C.E. and Lee, F.S. (2015) Treating the Developing versus Developed Brain: '
    'Translating Preclinical Mouse and Human Studies. Neuron, 86(6), pp. 1358-1368.',
    'Establishing neurobiological basis: adolescence as sensitive window for fear-learning and anxiety.',
    'Adolescent fear-learning circuitry is uniquely plastic and incomplete -- this developmental '
    'vulnerability is precisely what digital social environments exploit through social comparison '
    'and approval anxiety.',
    quality='High')

heading('3.2  Social Media as Anxiety Amplifier', level=2)
placeholder('Mechanisms: social comparison, approval anxiety, FOMO, cyberbullying')

citation('SM-2',
    'Fassi, L., Thomas, M., Parry, D.A. et al. (2024) Social Media Use and Internalizing Symptoms '
    'in Clinical and Community Adolescent Samples: A Systematic Review and Meta-Analysis. JAMA '
    'Pediatrics, 178(12), pp. 1213-1222.',
    'Core evidence: SM associated with internalising symptoms including anxiety (r=0.08-0.12, n>1.09M).',
    'The largest available meta-analysis confirms a consistent positive association between social '
    'media use and anxiety in adolescents -- with problematic use showing stronger effects.',
    quality='High')

citation('CB-1',
    'Li, Q., Liu, C., Knoll, N. and Obsuth, I. (2025) Mechanisms Linking Cyberbullying Victimisation '
    'to Internalising Problems in Youth: A Systematic Review and Meta-Analytic Structural Equation '
    'Modelling. Clinical Psychology Review, 116, pp. 102672.',
    'Proving mechanism: cyberbullying -> impaired emotion regulation -> anxiety (n=260,608, Meta-SEM).',
    'Cyberbullying operates through specific mediating mechanisms -- impaired emotion regulation and '
    'reduced social support -- providing clear, evidence-based intervention targets.',
    quality='High')

citation('DS-1',
    'Khetawat, D. and Steele, R.G. (2023) Examining the Association Between Digital Stress Components '
    'and Psychological Wellbeing: A Meta-Analysis. Clinical Child and Family Psychology Review, 26(4), '
    'pp. 1022-1037.',
    'Differentiating digital stress: approval anxiety and FOMO carry strongest anxiety associations (r=.34-.35).',
    'Digital stress is not monolithic -- approval anxiety and FOMO are the most anxiety-relevant '
    'components, identifying social comparison as the core mechanism for intervention.',
    quality='Medium-High')

heading('3.3  Balancing the Evidence', level=2)

citation('ST-1',
    'Odgers, C.L. and Jensen, M.R. (2020) Annual Research Review: Adolescent Mental Health in the '
    'Digital Age: Facts, Fears, and Future Directions. Journal of Child Psychology and Psychiatry, '
    '61(3), pp. 336-348.',
    'Critical counterpoint: pre-registered large-scale studies show small, often non-significant longitudinal effects.',
    'The moral panic around social media and adolescent mental health is not fully supported by '
    'methodologically rigorous research -- effects are real but modest, necessitating nuanced '
    'targeted intervention rather than blanket restrictions.',
    quality='High')

heading('3.4  Post-COVID Context', level=2)

citation('COVID-1',
    'Madigan, S., Racine, N., Vaillancourt, T. et al. (2023) Changes in Depression and Anxiety Among '
    'Children and Adolescents From Before to During the COVID-19 Pandemic: A Systematic Review and '
    'Meta-Analysis. JAMA Pediatrics, 177(6), pp. 567-581.',
    'Quantifying pandemic anxiety increase: SMC=0.10 across 53 longitudinal cohorts (n=40,807).',
    'COVID-19 produced measurable increases in adolescent anxiety, creating a heightened-need context '
    'in which both family-based and digital-aware interventions are more urgently required.',
    quality='High')

placeholder('EBSA doubled post-COVID -- link to M4 W4 P51; schools as entry point for identification')

citation('ICBT-1',
    'Nordh, M., Wahlund, T., Jolstedt, M. et al. (2021) Therapist-Guided Internet-Delivered Cognitive '
    'Behavioural Therapy vs Internet-Delivered Supportive Therapy for Children and Adolescents with '
    'Social Anxiety Disorder: A Randomized Clinical Trial. JAMA Psychiatry, 78(7), pp. 705-713.',
    'Internet-delivered CBT outperforms active control for social anxiety (d=0.67, RCT).',
    'Therapist-guided internet CBT meets adolescents in the digital environment where their anxiety '
    'is generated -- a key integration point between the societal and treatment dimensions.',
    quality='High')

divider()

# ---- 4. INTEGRATION ----
heading('4. Integration: A Systemic Response  (~300 words)', level=1)
placeholder('Family x digital interaction: anxious parents + unmonitored social media = compounding risk (gap in literature -- original argument)')
placeholder('Continuum of need as organising framework (M4 W3 P35 curriculum)')
placeholder('Stepped care: universal (digital literacy, SEL) > targeted (SPACE, family CBT) > specialist (CBT + SSRI)')

citation('CBT-1',
    'Sigurvinsdottir, A.L., Jensinudottir, K.B., Baldvinsdottir, K.D. et al. (2019) Effectiveness of '
    'Cognitive Behavioural Therapy for Child and Adolescent Anxiety Disorders Across Different CBT '
    'Modalities and Comparisons: A Systematic Review and Meta-Analysis. Nordic Journal of Psychiatry, '
    '73(3), pp. 168-180.',
    'All CBT modalities (individual, family, group, internet) are efficacious -- supporting stepped-care logic.',
    'The equivalence of CBT modalities supports deploying less intensive formats (digital, group) '
    'first, with specialist individual CBT reserved for non-responders -- a resource-efficient '
    'approach to the undertreatment gap.',
    quality='High')

placeholder('Evidence gaps: no RCT of full stepped-care model for child anxiety; SM-anxiety causality contested; FA research concentrated in Lebowitz group')
placeholder('Future directions: family-digital interaction research; longitudinal post-COVID data')
divider()

# ---- 5. CONCLUSION ----
heading('5. Conclusion  (~200 words)', level=1)
placeholder('Restate thesis: anxiety is systemic -- family dynamics + digital environment, not individual pathology')
placeholder('Summary: family accommodation and parental transmission --> family-inclusive CBT (SPACE); neurobiological vulnerability + SM mechanisms --> digital-aware, therapist-guided intervention')
placeholder('Integrative vision: stepped-care across family, school, and digital contexts')
placeholder('Closing: individual-treatment model is insufficient; evidence demands a systemic response')

doc.add_page_break()

# ---- REFERENCE LIST ----
heading('Reference List', level=1)
p = doc.add_paragraph()
sf(p.add_run('(Alphabetical -- Harvard WMS format. Full list to be audited in Phase 8 FQA.)'), 9, italic=True, color=(140, 140, 140))

refs = [
    'Ahmadzadeh, Y., Schoeler, T., Han, M. et al. (2021) Systematic review and meta-analysis of genetically informed research: Associations between parent anxiety and offspring internalizing problems. Journal of the American Academy of Child and Adolescent Psychiatry, 60(8), pp. 959-972.',
    'Ahmed, M., Walsh, E., Dawel, A. et al. (2024) Social media use, mental health and sleep: A systematic review with meta-analyses. Journal of Affective Disorders, 362, pp. 217-228.',
    'Bennett, S., McClure, G. and Yap, M. (2025) Differences in parental factors between parents with and without depression or anxiety: A systematic review and meta-analysis. Journal of Affective Disorders, 369, pp. 120085.',
    'Casey, B.J., Glatt, C.E. and Lee, F.S. (2015) Treating the developing versus developed brain: Translating preclinical mouse and human studies. Neuron, 86(6), pp. 1358-1368.',
    'Fassi, L., Thomas, M., Parry, D.A. et al. (2024) Social media use and internalizing symptoms in clinical and community adolescent samples: A systematic review and meta-analysis. JAMA Pediatrics, 178(12), pp. 1213-1222.',
    'Fox, R. and Fleming, L. (2025) Parental cognitions and child anxiety: A systematic review. Journal of Anxiety Disorders, 101, pp. 103021.',
    'Khetawat, D. and Steele, R.G. (2023) Examining the association between digital stress components and psychological wellbeing: A meta-analysis. Clinical Child and Family Psychology Review, 26(4), pp. 1022-1037.',
    'Lebowitz, E.R., Marin, C., Martino, A. et al. (2019) Parent-based treatment as efficacious as cognitive behavioral therapy for childhood anxiety: A randomized noninferiority study of SPACE. Journal of the American Academy of Child and Adolescent Psychiatry, 58(12), pp. 1170-1181.',
    'Li, Q., Liu, C., Knoll, N. and Obsuth, I. (2025) Mechanisms linking cyberbullying victimisation to internalising problems in youth: A systematic review and meta-analytic structural equation modelling. Clinical Psychology Review, 116, pp. 102672.',
    'Loades, M.E., Chatburn, E., Higson-Sweeney, N. et al. (2020) Rapid systematic review: The impact of social isolation and loneliness on the mental health of children and adolescents in the context of COVID-19. Journal of the American Academy of Child and Adolescent Psychiatry, 59(11), pp. 1218-1239.',
    'Madigan, S., Racine, N., Vaillancourt, T. et al. (2023) Changes in depression and anxiety among children and adolescents from before to during the COVID-19 pandemic: A systematic review and meta-analysis. JAMA Pediatrics, 177(6), pp. 567-581.',
    'Nordh, M., Wahlund, T., Jolstedt, M. et al. (2021) Therapist-guided internet-delivered cognitive behavioural therapy vs internet-delivered supportive therapy for children and adolescents with social anxiety disorder: A randomized clinical trial. JAMA Psychiatry, 78(7), pp. 705-713.',
    'Odgers, C.L. and Jensen, M.R. (2020) Annual research review: Adolescent mental health in the digital age: Facts, fears, and future directions. Journal of Child Psychology and Psychiatry, 61(3), pp. 336-348.',
    'Sigurvinsdottir, A.L., Jensinudottir, K.B., Baldvinsdottir, K.D. et al. (2019) Effectiveness of cognitive behavioural therapy for child and adolescent anxiety disorders across different CBT modalities and comparisons: A systematic review and meta-analysis. Nordic Journal of Psychiatry, 73(3), pp. 168-180.',
    'Wang, Z., Whiteside, S.P.H., Sim, L. et al. (2017) Comparative effectiveness and safety of cognitive behavioral therapy and pharmacotherapy for childhood anxiety disorders: A systematic review and meta-analysis. JAMA Pediatrics, 171(11), pp. 1049-1056.',
    'Zecchinato, F., Ahmadzadeh, Y., Kreppner, J. and Lawrence, P.J. (2024) Paternal anxiety and emotional and behavioral outcomes in offspring: A systematic review and meta-analysis. Journal of the American Academy of Child and Adolescent Psychiatry, 63(11), pp. 1142-1156.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(4)
    sf(p.add_run(ref), 9.5)

out = 'C:/Users/tango/OneDrive/10 Documents/15 Dev/EssayFabrik/01 projects/module4/06_research/essay_draft_scaffold.docx'
doc.save(out)
print('Saved:', out)
