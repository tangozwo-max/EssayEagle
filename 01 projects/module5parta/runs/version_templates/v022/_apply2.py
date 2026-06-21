# -*- coding: utf-8 -*-
# v022 round 2: engagement fix + lossless word-count trim. Preserves prior highlights.
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

PATH = r"C:\Users\tango\OneDrive\10 Documents\15 Dev\EssayFabrik\01 projects\module4\runs\v022\v022.docx"
doc = Document(PATH)

YS = chr(0xE000); YE = chr(0xE001); GS = chr(0xE002); GE = chr(0xE003)

def find_para(sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    raise SystemExit("PARA NOT FOUND: " + sub[:50])

def segment(text):
    segs=[]; buf=''; color=None
    for c in text:
        if c == YS or c == GS:
            if buf: segs.append((buf,color)); buf=''
            color = 'Y' if c == YS else 'G'
        elif c == YE or c == GE:
            if buf: segs.append((buf,color)); buf=''
            color = None
        else:
            buf += c
    if buf: segs.append((buf,color))
    return segs

def rebuild(para, segs):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for txt,color in segs:
        if not txt: continue
        run = para.add_run(txt)
        if color == 'Y': run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif color == 'G': run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

def must_replace(t, old, new):
    if old not in t:
        raise SystemExit("OLD NOT FOUND: " + old[:60])
    return t.replace(old, new, 1)

# 1) endorsement -> engagement (yellow). Isolated paragraph, no prior highlights.
p = find_para("reduce endorsement in the child")
t = must_replace(p.text, "reduce endorsement in the child",
                 "reduce " + YS + "engagement" + YE + " in the child")
rebuild(p, segment(t))

# 2) R3 paragraph: re-mark prior green spans, trim forms tail, trim 'go on to' (yellow)
p = find_para("Its main forms are generalised")
t = p.text
t = must_replace(t, "(Beesdo et al., 2009)", GS + "(Beesdo et al., 2009)" + GE)
t = must_replace(t,
    "Its main forms are generalised, separation and social anxiety and specific phobias, each taking age-specific forms.",
    GS + "Its main forms are generalised, separation and social anxiety and specific phobias." + GE)
t = must_replace(t, "seem to go on to face", YS + "seem to face" + YE)
rebuild(p, segment(t))

# 3) R8 single-session definition: trim, keep green
p = find_para("single-session intervention")
t = p.text
t = must_replace(t,
    "a one-off, self-contained session designed to help in a single contact,",
    GS + "a self-contained, one-off session," + GE)
rebuild(p, segment(t))

doc.save(PATH)
print("SAVED OK")
