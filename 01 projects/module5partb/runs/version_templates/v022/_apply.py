# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

PATH = r"C:\Users\tango\OneDrive\10 Documents\15 Dev\EssayFabrik\01 projects\module4\runs\v022\v022.docx"
doc = Document(PATH)

YS = chr(0xE000); YE = chr(0xE001); GS = chr(0xE002); GE = chr(0xE003)

def find_para(sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    raise SystemExit("PARA NOT FOUND: " + sub[:50])

def segment(text):
    segs=[]; buf=''; color=None
    for c in text:
        if c == YS or c == GS:
            if buf: segs.append((buf,color)); buf=''
            color = 'Y' if c == YS else 'G'
        elif c == YE or c == GE:
            if buf: segs.append((buf,color)); buf=''
            color = None
        else:
            buf += c
    if buf: segs.append((buf,color))
    return segs

def rebuild(para, segs):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for txt,color in segs:
        if not txt: continue
        run = para.add_run(txt)
        if color == 'Y': run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif color == 'G': run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

def edit(sub, edits):
    p = find_para(sub)
    text = p.text
    for old,new,color in edits:
        if old not in text:
            raise SystemExit("OLD NOT FOUND in para [%s]: %r" % (sub[:30], old))
        if color is None:
            rep = new
        else:
            s,e = (YS,YE) if color=='Y' else (GS,GE)
            rep = s + new + e
        text = text.replace(old, rep, 1)
    rebuild(p, segment(text))

def replace_whole(sub, new, color):
    rebuild(find_para(sub), [(new,color)])

def insert_ref(anchor, text):
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Bibliography' and p.text.startswith(anchor):
            newp = p.insert_paragraph_before(text, 'Bibliography')
            for run in newp.runs:
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            return
    raise SystemExit("ANCHOR NOT FOUND: " + anchor)

# R1
edit("most common psychiatric condition of childhood",
     [("(SOURCE)","(Polanczyk et al., 2015)","G")])
# R2 + R3
edit("developmentally ordinary fear",
     [("(SOURCE)","(Beesdo et al., 2009)","G"),
      ("It takes age-specific forms.",
       "Its main forms are generalised, separation and social anxiety and specific phobias, each taking age-specific forms.","G")])
# R5
edit("Working in a context of case formulation",
     [("perpetuating (continuing) and precipitating (triggering) factors of anxiety",
       "perpetuating (maintaining) factors of anxiety","Y"),
      ("rather than the fear itself",
       "rather than the precipitating (triggering) events or the fear itself","Y")])
# R6 + R7
edit("When CBT alone is not enough",
     [("(SOURCE)","(Wang et al., 2017)","G"),
      ("So in practice, SSRIs are reserved for more severe cases or children who have not responded to first-line CBT, with careful monitoring (SOURCE).",
       "So SSRIs should be reserved for more severe cases or children who have not responded to first-line CBT, with careful monitoring.","Y")])
# R8
edit("A leaner option, the single-session intervention",
     [("single-session intervention, reaches young people",
       "single-session intervention, " + GS + "a one-off, self-contained session designed to help in a single contact, " + GE + "reaches young people",
       None)])
# R9
edit("Young children regulate emotion through a caregiver",
     [("(SOURCE)","(Morris et al., 2007)","G")])
# R10
old_frontier=("A newer technological frontier attempts to address both, direct acess. "
 "Conversational-agent chatbots deliver automated support, and virtual-reality exposure lets a child face feared situations in a controlled, simulated setting. "
 "Both target the reach-and-engagement gap for a digital generation, and both are acceptable to young people, but neither has yet been tested against established therapist-led treatment, so neither can yet count as part of the answer (Park et al., 2025; Baschab et al., 2026).")
new_frontier=("A newer frontier attempts this with technology. "
 "Conversational-agent chatbots offer a low-barrier, anonymous route that could reach young people directly, including those a parent does not bring forward (Park et al., 2025), while virtual-reality exposure adds immersion, letting a child face feared situations in a controlled, simulated setting (Baschab et al., 2026). "
 "Both are acceptable to young people, but neither has yet been tested against established therapist-led treatment, so neither can yet count as part of the answer.")
edit("A newer technological frontier", [(old_frontier,new_frontier,"Y")])
# R12
new_concl=("Our interventions have kept pace as treatments, but not as a system. "
 "Cognitive behavioural therapy, parent-led work and the careful use of medication help children, and they do so in ordinary clinics, not only in trials. "
 "The shortfall lies elsewhere. Most children who need care still do not reach it, the youngest depend on adults who cannot always act, and new cases keep arising from a digital environment that treatment alone cannot touch. "
 "The gap is therefore one of delivery and prevention, not of therapies. "
 "Closing it calls less for a new treatment than for delivering the ones we have more widely, through parents and guided-digital routes, and for targeting prevention where risk is concentrated rather than spreading it thinly. "
 "The evidence for what works already exists. The task now is to reach the children who need it, many of whom still depend on the adults around them to get there.")
replace_whole("kept pace only in part", new_concl, "Y")
# new references
insert_ref("Bie, F.", "Beesdo, K., Knappe, S., & Pine, D.S. (2009) Anxiety and anxiety disorders in children and adolescents: developmental issues and implications for DSM-V. Psychiatric Clinics of North America, 32 (3), 483–524. Available from: https://doi.org/10.1016/j.psc.2009.06.002 (Accessed 12 June 2026).")
insert_ref("Nordh, M.", "Morris, A.S., Silk, J.S., Steinberg, L., Myers, S.S., & Robinson, L.R. (2007) The role of the family context in the development of emotion regulation. Social Development, 16 (2), 361–388. Available from: https://doi.org/10.1111/j.1467-9507.2007.00389.x (Accessed 12 June 2026).")
insert_ref("Radez, J.", "Polanczyk, G.V., Salum, G.A., Sugaya, L.S., Caye, A., & Rohde, L.A. (2015) Annual Research Review: A meta-analysis of the worldwide prevalence of mental disorders in children and adolescents. Journal of Child Psychology and Psychiatry, 56 (3), 345–365. Available from: https://doi.org/10.1111/jcpp.12381 (Accessed 12 June 2026).")

doc.save(PATH)
print("SAVED OK")
